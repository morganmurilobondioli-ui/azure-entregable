from pathlib import Path
import shutil
import subprocess
import textwrap
import zipfile

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from lxml import etree
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
TEMPLATE = Path(r"C:\Users\morga\Downloads\PIAD-627_FORMATOALUMNOTRABAJOFINAL.docx")
OUT_DOCX = ROOT / "entregables" / "PIAD-627_TRABAJO_FINAL_INNOVVENTAS.docx"
OUT_MD = ROOT / "docs" / "propuesta_innovventas.md"
ASSET_DIR = ROOT / "assets"


AZURE_SOURCES = [
    (
        "Microsoft Copilot Studio documentation",
        "https://learn.microsoft.com/en-us/microsoft-copilot-studio/",
    ),
    (
        "Explore AI capabilities in Copilot Studio",
        "https://learn.microsoft.com/en-us/microsoft-copilot-studio/guidance/ai-capabilities",
    ),
    (
        "Web Chat overview - Azure AI Bot Service",
        "https://learn.microsoft.com/en-us/azure/bot-service/bot-builder-webchat-overview?view=azure-bot-service-4.0",
    ),
    (
        "About Direct Line - Azure Bot Service",
        "https://learn.microsoft.com/en-gb/azure/bot-service/bot-service-channel-directline?view=azure-bot-service-4.0",
    ),
    (
        "Add telemetry to your bot - Azure Bot Service",
        "https://learn.microsoft.com/en-us/azure/bot-service/bot-builder-telemetry?view=azure-bot-service-4.0",
    ),
]


FAQS = [
    "Disponibilidad de laptops, celulares, accesorios y componentes.",
    "Especificaciones tecnicas: procesador, memoria, almacenamiento, garantia y compatibilidad.",
    "Metodos de pago aceptados, cuotas, validacion de pago y comprobante.",
    "Costos y tiempos de envio, cobertura por zona y retiro en tienda.",
    "Estado del pedido, cambios, devoluciones y politicas de garantia.",
    "Recomendacion de producto segun presupuesto y uso: estudios, gaming, oficina o diseno.",
    "Soporte tecnico basico: configuracion inicial, fallas comunes y contacto con soporte humano.",
]


GUIDE_QA = [
    (
        "Pregunta 01: ¿Cuales son las preguntas frecuentes mas comunes que enfrentan los clientes al usar el sitio web de InnovVentas?",
        "Las consultas mas frecuentes se concentran en disponibilidad de stock, especificaciones tecnicas, metodos de pago, tiempos de envio, garantias, cambios/devoluciones, seguimiento de pedidos y soporte tecnico basico. Tambien aparecen preguntas comparativas como 'que laptop me conviene para estudiar' o 'que celular tiene mejor bateria'. Estas consultas deben priorizarse porque ocurren justo antes de agregar al carrito o completar el pago, momento donde la falta de respuesta inmediata genera abandono.",
    ),
    (
        "Pregunta 02: ¿Que herramientas o plataformas son mas adecuadas para desarrollar e integrar el chatbot al sitio web?",
        "La solucion recomendada es Microsoft Copilot Studio para construir el agente con bajo codigo, temas guiados, respuestas generativas basadas en fuentes de conocimiento y conexion con acciones. Para una integracion web mas personalizada se propone Azure AI Bot Service con Bot Framework Web Chat y canal Direct Line. Application Insights/Azure Monitor permitira registrar disponibilidad, uso, errores e interacciones del bot. Esta combinacion aprovecha servicios Azure, reduce tiempo de implementacion y permite evolucionar desde FAQ hacia asistencia transaccional.",
    ),
    (
        "Pregunta 03: ¿Como se puede evaluar la efectividad del chatbot en terminos de satisfaccion del cliente y aumento de ventas?",
        "La efectividad debe medirse comparando una linea base previa contra indicadores posteriores a la implementacion. Para satisfaccion: CSAT al cierre de la conversacion, porcentaje de conversaciones resueltas sin agente humano, tiempo medio de respuesta y tasa de escalamiento. Para ventas: conversion de usuarios que interactuaron con el chatbot, reduccion de carritos abandonados, productos agregados al carrito luego de una recomendacion y valor promedio de pedido. Se recomienda revisar resultados semanalmente durante el piloto y mensualmente en produccion.",
    ),
    (
        "Pregunta 04: ¿Que desafios tecnicos podrian surgir durante la implementacion del chatbot y como podrian resolverse?",
        "Los principales desafios son mantener actualizado el catalogo, proteger datos personales, integrar el bot con sistemas de stock/pedidos, evitar respuestas imprecisas y medir correctamente el impacto comercial. Para resolverlos se propone conectar el bot a APIs o base de datos de productos, aplicar autenticacion solo cuando sea necesaria, definir temas criticos con respuestas controladas, usar fuentes de conocimiento verificadas, activar escalamiento a soporte humano y monitorear errores con Application Insights.",
    ),
    (
        "Pregunta 05: ¿Que metricas deben monitorearse para asegurar el exito del chatbot y como se pueden optimizar sus funcionalidades?",
        "Deben monitorearse sesiones iniciadas, intenciones mas consultadas, tasa de resolucion, preguntas sin respuesta, tiempo medio de respuesta, CSAT, escalamiento, errores tecnicos, conversion asistida, abandono de carrito y ventas atribuidas. La optimizacion se realiza revisando transcripciones, creando nuevos temas para preguntas recurrentes, mejorando la base de conocimiento, ajustando mensajes de recomendacion, probando variantes de flujos y priorizando las intenciones con mayor impacto en conversion.",
    ),
]


PROPOSAL_SECTIONS = [
    (
        "1. Necesidad del cliente y problema identificado",
        [
            "InnovVentas pierde oportunidades de venta porque el sitio web no responde de forma inmediata dudas que aparecen durante la decision de compra. El problema afecta tres puntos del flujo e-commerce: exploracion de productos, confirmacion de condiciones de compra y soporte posterior.",
            "La propuesta consiste en implementar un chatbot integrado al sitio web, disponible 24/7, capaz de resolver preguntas frecuentes, guiar la seleccion de productos, reducir friccion en checkout y derivar a soporte humano cuando la consulta supere el alcance automatizado.",
        ],
    ),
    (
        "2. Alcance funcional del chatbot",
        [
            "Atendera preguntas frecuentes sobre stock, especificaciones, pagos, envios, garantias y devoluciones.",
            "Recomendara productos segun categoria, presupuesto, uso previsto y caracteristicas relevantes.",
            "Acompanara el proceso de compra con enlaces a producto, carrito, medios de pago y seguimiento de pedido.",
            "Ofrecera soporte tecnico basico y escalara a un asesor cuando detecte reclamos, fallas complejas o baja satisfaccion.",
        ],
    ),
    (
        "3. Plataforma seleccionada",
        [
            "Se recomienda Microsoft Copilot Studio como plataforma principal por su enfoque de bajo codigo, capacidad de crear agentes, configurar temas, usar conocimiento corporativo y publicar el agente en canales web. Para un sitio con necesidades de mayor personalizacion visual o control del front-end, se complementa con Azure AI Bot Service, Bot Framework Web Chat y Direct Line.",
            "Application Insights/Azure Monitor se usara para telemetria, diagnostico y analisis de uso. Las fuentes oficiales de Microsoft indican que Web Chat puede integrarse en sitios web mediante JavaScript o React y que Application Insights permite visualizar disponibilidad, rendimiento y uso del bot.",
        ],
    ),
    (
        "4. Plan de implementacion",
        [
            "Fase 1 - Diagnostico: analizar transcripciones, formularios, busquedas internas y consultas de ventas para construir una matriz de preguntas frecuentes.",
            "Fase 2 - Diseno: definir intenciones, entidades, respuestas, politicas de escalamiento y tono de comunicacion.",
            "Fase 3 - Construccion: crear el agente en Copilot Studio, cargar fuentes de conocimiento, configurar temas y acciones con APIs de catalogo/pedidos.",
            "Fase 4 - Integracion: publicar el chatbot en el sitio web con Web Chat/Direct Line o canal web de Copilot Studio, respetando identidad visual y accesibilidad.",
            "Fase 5 - Pruebas: validar respuestas, casos limite, seguridad, privacidad, tiempos de respuesta y trazabilidad de eventos.",
            "Fase 6 - Piloto y mejora: desplegar a un porcentaje de usuarios, medir metricas y optimizar temas antes del despliegue total.",
        ],
    ),
    (
        "5. Seguridad, calidad y medio ambiente",
        [
            "Seguridad: minimizar datos personales, evitar exponer claves de Direct Line en el cliente, usar tokens generados por backend, aplicar HTTPS, control de acceso y retencion responsable de transcripciones.",
            "Calidad: probar preguntas frecuentes, variantes de lenguaje, disponibilidad de integraciones, respuestas sin informacion suficiente y experiencia movil.",
            "Medio ambiente: usar servicios cloud escalables bajo demanda, evitar infraestructura sobredimensionada y reutilizar documentacion digital para capacitacion y mejora continua.",
        ],
    ),
    (
        "6. Indicadores de exito",
        [
            "Reducir en al menos 15% el abandono de carrito durante los primeros tres meses.",
            "Resolver automaticamente al menos 60% de preguntas frecuentes del sitio.",
            "Mantener CSAT igual o superior a 4/5 en conversaciones cerradas.",
            "Detectar semanalmente nuevas preguntas sin respuesta y convertirlas en mejoras del flujo.",
        ],
    ),
]


STEPS = [
    ("1. Levantar informacion del caso y analizar el flujo e-commerce.", "Documentar fuentes de informacion, respetar privacidad de clientes y usar datos anonimizados."),
    ("2. Identificar preguntas frecuentes y agruparlas por intencion.", "Aplicar criterios de calidad de datos y validar con ventas/soporte."),
    ("3. Disenar mapa conversacional y reglas de escalamiento.", "Incluir mensajes claros, accesibles y sin promesas comerciales no verificadas."),
    ("4. Seleccionar plataforma: Copilot Studio + Azure Bot Service/Web Chat.", "Usar documentacion oficial Microsoft y servicios compatibles con Azure."),
    ("5. Crear base de conocimiento de productos, pagos, envios y garantias.", "Mantener contenido actualizado, versionado y aprobado por el negocio."),
    ("6. Construir temas del chatbot y respuestas generativas controladas.", "Aplicar principios de IA responsable, revision humana y fuentes confiables."),
    ("7. Integrar con catalogo, stock, pedidos y soporte humano.", "Proteger credenciales, usar HTTPS y controlar permisos de API."),
    ("8. Publicar el widget en el sitio web.", "No exponer secretos en front-end; generar tokens desde backend cuando use Direct Line."),
    ("9. Configurar telemetria en Application Insights/Azure Monitor.", "Medir disponibilidad, rendimiento, uso, errores e indicadores de negocio."),
    ("10. Ejecutar pruebas funcionales, seguridad y experiencia de usuario.", "Validar accesibilidad, navegacion movil, carga rapida y manejo de errores."),
    ("11. Lanzar piloto controlado y comparar contra linea base.", "Monitorear resultados sin manipular datos personales innecesarios."),
    ("12. Optimizar continuamente a partir de metricas y transcripciones.", "Gestionar cambios documentados, reducir desperdicio operativo y mejorar calidad."),
]


def set_cell_text(cell, text, bold=False, size=9):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_cell_paragraph(cell, text, bold=False, size=9, space_after=4):
    p = cell.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    return p


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def format_doc(doc):
    for section in doc.sections:
        section.top_margin = Inches(0.65)
        section.bottom_margin = Inches(0.65)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
    for p in doc.paragraphs:
        for run in p.runs:
            run.font.name = "Arial"
            if run.font.size is None:
                run.font.size = Pt(10)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.name = "Arial"
                        if run.font.size is None:
                            run.font.size = Pt(9)


def load_font(size=24, bold=False):
    candidates = [
        r"C:\Windows\Fonts\arialbd.ttf" if bold else r"C:\Windows\Fonts\arial.ttf",
        r"C:\Windows\Fonts\calibrib.ttf" if bold else r"C:\Windows\Fonts\calibri.ttf",
    ]
    for candidate in candidates:
        if Path(candidate).exists():
            return ImageFont.truetype(candidate, size)
    return ImageFont.load_default()


def wrap_lines(draw, text, font, width):
    words = text.split()
    lines = []
    current = ""
    for word in words:
        trial = (current + " " + word).strip()
        bbox = draw.textbbox((0, 0), trial, font=font)
        if bbox[2] <= width or not current:
            current = trial
        else:
            lines.append(current)
            current = word
    if current:
        lines.append(current)
    return lines


def draw_box(draw, xy, text, fill, outline="#1f2937"):
    x1, y1, x2, y2 = xy
    draw.rounded_rectangle(xy, radius=18, fill=fill, outline=outline, width=3)
    font = load_font(25, bold=True)
    small = load_font(20)
    parts = text.split("\n", 1)
    title = parts[0]
    body = parts[1] if len(parts) > 1 else ""
    y = y1 + 18
    for line in wrap_lines(draw, title, font, x2 - x1 - 28):
        bbox = draw.textbbox((0, 0), line, font=font)
        draw.text((x1 + (x2 - x1 - bbox[2]) / 2, y), line, fill="#111827", font=font)
        y += 31
    if body:
        y += 6
        for line in wrap_lines(draw, body, small, x2 - x1 - 28)[:3]:
            bbox = draw.textbbox((0, 0), line, font=small)
            draw.text((x1 + (x2 - x1 - bbox[2]) / 2, y), line, fill="#374151", font=small)
            y += 25


def arrow(draw, start, end):
    draw.line([start, end], fill="#374151", width=4)
    x1, y1 = start
    x2, y2 = end
    if abs(x2 - x1) >= abs(y2 - y1):
        direction = 1 if x2 >= x1 else -1
        points = [(x2, y2), (x2 - 16 * direction, y2 - 9), (x2 - 16 * direction, y2 + 9)]
    else:
        direction = 1 if y2 >= y1 else -1
        points = [(x2, y2), (x2 - 9, y2 - 16 * direction), (x2 + 9, y2 - 16 * direction)]
    draw.polygon(points, fill="#374151")


def make_diagrams():
    ASSET_DIR.mkdir(exist_ok=True)

    conv = Image.new("RGB", (1500, 920), "#ffffff")
    d = ImageDraw.Draw(conv)
    title_font = load_font(36, bold=True)
    d.text((50, 35), "Mapa conversacional propuesto", fill="#0f172a", font=title_font)
    boxes = {
        "start": (80, 120, 380, 230, "Inicio\nSaludo y opciones", "#dbeafe"),
        "intent": (600, 120, 900, 230, "Detectar intencion\nFAQ, compra o soporte", "#dcfce7"),
        "faq": (80, 360, 380, 505, "FAQ\nStock, pagos, envios, garantia", "#fef3c7"),
        "purchase": (600, 360, 900, 505, "Asistente de compra\nRecomienda y guia carrito", "#ede9fe"),
        "support": (1120, 360, 1420, 505, "Soporte tecnico\nDiagnostico basico", "#fee2e2"),
        "resolve": (340, 655, 640, 790, "Resolver\nRespuesta + enlace util", "#ccfbf1"),
        "human": (860, 655, 1160, 790, "Escalar\nAsesor humano si aplica", "#e5e7eb"),
    }
    for key, (x1, y1, x2, y2, text, fill) in boxes.items():
        draw_box(d, (x1, y1, x2, y2), text, fill)
    arrow(d, (380, 175), (600, 175))
    arrow(d, (690, 230), (250, 360))
    arrow(d, (750, 230), (750, 360))
    arrow(d, (860, 230), (1270, 360))
    arrow(d, (230, 505), (430, 655))
    arrow(d, (750, 505), (520, 655))
    arrow(d, (1270, 505), (1010, 655))
    d.text((50, 845), "Cierre: registrar metrica, pedir CSAT y actualizar base de conocimiento.", fill="#334155", font=load_font(24))
    conv.save(ASSET_DIR / "mapa_conversacional.png")

    arch = Image.new("RGB", (1500, 860), "#ffffff")
    d = ImageDraw.Draw(arch)
    d.text((50, 35), "Arquitectura de solucion Azure", fill="#0f172a", font=title_font)
    boxes = [
        (70, 140, 350, 255, "Cliente web\nPC o movil", "#dbeafe"),
        (460, 140, 740, 255, "Sitio e-commerce\nWidget Web Chat", "#e0f2fe"),
        (850, 140, 1130, 255, "Direct Line\nCanal seguro", "#dcfce7"),
        (460, 375, 740, 505, "Copilot Studio / Bot Service\nTemas, IA y dialogos", "#ede9fe"),
        (70, 610, 350, 725, "Catalogo y stock\nAPI productos", "#fef3c7"),
        (460, 610, 740, 725, "Pedidos y pagos\nAPI e-commerce", "#fee2e2"),
        (850, 610, 1130, 725, "Application Insights\nMetricas y errores", "#e5e7eb"),
        (1210, 375, 1460, 505, "Soporte humano\nEscalamiento", "#ccfbf1"),
    ]
    for x1, y1, x2, y2, text, fill in boxes:
        draw_box(d, (x1, y1, x2, y2), text, fill)
    arrow(d, (350, 197), (460, 197))
    arrow(d, (740, 197), (850, 197))
    arrow(d, (990, 255), (600, 375))
    arrow(d, (600, 505), (230, 610))
    arrow(d, (600, 505), (600, 610))
    arrow(d, (740, 450), (850, 665))
    arrow(d, (740, 440), (1210, 440))
    d.text((50, 790), "Principio clave: no exponer secretos en el navegador; generar tokens desde backend y monitorear eventos de negocio.", fill="#334155", font=load_font(23))
    arch.save(ASSET_DIR / "arquitectura_azure.png")


def fill_template():
    if not TEMPLATE.exists():
        raise FileNotFoundError(f"No existe la plantilla: {TEMPLATE}")
    OUT_DOCX.parent.mkdir(exist_ok=True)
    shutil.copyfile(TEMPLATE, OUT_DOCX)
    doc = Document(str(OUT_DOCX))
    format_doc(doc)

    # Datos del estudiante.
    t = doc.tables[0]
    set_cell_text(t.cell(0, 1), "[Completar apellidos y nombres]", size=8)
    set_cell_text(t.cell(0, 3), "[Completar ID]", size=8)
    set_cell_text(t.cell(1, 1), "[Completar Direccion Zonal/CFP]", size=8)
    set_cell_text(t.cell(2, 1), "Tecnologias de la Informacion", size=8)
    set_cell_text(t.cell(2, 3), "[Completar semestre]", size=8)
    set_cell_text(t.cell(3, 1), "AI-900T00 Conceptos Basicos de IA en Microsoft Azure", size=8)
    set_cell_text(t.cell(4, 1), "Uso de chatbots en E-commerce", size=8)

    # Preguntas guia.
    t = doc.tables[1]
    for idx, (question, answer) in enumerate(GUIDE_QA):
        set_cell_text(t.cell(idx * 2, 0), question, bold=True, size=8)
        set_cell_text(t.cell(idx * 2, 1), "", size=8)
        set_cell_text(t.cell(idx * 2 + 1, 0), answer, size=8)
        shade_cell(t.cell(idx * 2, 0), "E8F0FE")

    # Cronograma.
    t = doc.tables[2]
    headers = ["N", "ACTIVIDADES", "Semana 1", "Semana 2", "Semana 3", "Semana 4", "Semana 5", "Semana 6"]
    for col, value in enumerate(headers):
        set_cell_text(t.cell(1, col), value, bold=True, size=8)
        shade_cell(t.cell(1, col), "D9EAF7")
    activities = [
        ("1", "Analisis del caso y preguntas frecuentes", ["X", "", "", "", "", ""]),
        ("2", "Diseno de flujos conversacionales", ["", "X", "", "", "", ""]),
        ("3", "Seleccion y configuracion de plataforma Azure", ["", "X", "X", "", "", ""]),
        ("4", "Construccion de temas, FAQ e integraciones", ["", "", "X", "X", "", ""]),
        ("5", "Pruebas funcionales, seguridad y metricas", ["", "", "", "X", "X", ""]),
        ("6", "Piloto, optimizacion y entrega final", ["", "", "", "", "X", "X"]),
    ]
    for row_idx, (n, act, marks) in enumerate(activities, start=2):
        set_cell_text(t.cell(row_idx, 0), n, size=8)
        set_cell_text(t.cell(row_idx, 1), act, size=8)
        for i, mark in enumerate(marks, start=2):
            set_cell_text(t.cell(row_idx, i), mark, bold=bool(mark), size=8)

    # Recursos.
    resources = [
        (3, [("Computadora portatil o PC", "1"), ("Conexion a Internet", "1"), ("Cuenta Microsoft/Azure", "1"), ("Servidor o hosting del sitio web", "1")]),
        (4, [("Microsoft Copilot Studio", "1"), ("Azure AI Bot Service", "1"), ("Bot Framework Web Chat / Direct Line", "1"), ("Application Insights / Azure Monitor", "1")]),
        (5, [("Catalogo de productos", "1"), ("Base de preguntas frecuentes", "1"), ("Politicas de pago, envio y garantia", "1"), ("Transcripciones anonimizadas", "Segun disponibilidad")]),
    ]
    for table_index, items in resources:
        table = doc.tables[table_index]
        for r, (desc, qty) in enumerate(items, start=2):
            set_cell_text(table.cell(r, 0), desc, size=8)
            set_cell_text(table.cell(r, 1), qty, size=8)

    # Propuesta de solucion.
    t = doc.tables[6]
    cell = t.cell(1, 0)
    cell.text = ""
    add_cell_paragraph(cell, "Propuesta tecnica para InnovVentas", bold=True, size=12)
    for title, paragraphs in PROPOSAL_SECTIONS:
        add_cell_paragraph(cell, title, bold=True, size=10)
        for paragraph in paragraphs:
            add_cell_paragraph(cell, paragraph, size=8, space_after=3)
    add_cell_paragraph(cell, "Preguntas frecuentes priorizadas", bold=True, size=10)
    for item in FAQS:
        add_cell_paragraph(cell, item, size=8, space_after=2)
    add_cell_paragraph(cell, "Diagramas de la propuesta", bold=True, size=10)
    for image_name in ["mapa_conversacional.png", "arquitectura_azure.png"]:
        p = cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(ASSET_DIR / image_name), width=Inches(6.3))
    add_cell_paragraph(cell, "Fuentes consultadas", bold=True, size=10)
    for name, url in AZURE_SOURCES:
        add_cell_paragraph(cell, f"{name}: {url}", size=7, space_after=1)

    # Operaciones, normas y estandares.
    t = doc.tables[7]
    for row_idx, (op, norm) in enumerate(STEPS, start=1):
        set_cell_text(t.cell(row_idx, 0), op, size=8)
        set_cell_text(t.cell(row_idx, 1), norm, size=8)

    # Dibujo / esquema.
    t = doc.tables[8]
    diagram_cell = t.cell(0, 0)
    diagram_cell.text = ""
    add_cell_paragraph(diagram_cell, "Mapa conversacional e integracion Azure propuesta", bold=True, size=9)
    for image_name in ["mapa_conversacional.png", "arquitectura_azure.png"]:
        p = diagram_cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(ASSET_DIR / image_name), width=Inches(6.0))
    set_cell_text(t.cell(1, 1), "Uso de chatbots en E-commerce", bold=True, size=8)
    set_cell_text(t.cell(2, 1), "[Completar apellidos y nombres]", size=8)
    set_cell_text(t.cell(2, 2), "Escala referencial", size=8)

    # Evidencias.
    t = doc.tables[9]
    for row_idx in range(1, len(t.rows)):
        row_cells = t.rows[row_idx].cells
        if len(row_cells) >= 3:
            set_cell_text(row_cells[1], "X", bold=True, size=9)
            set_cell_text(row_cells[2], "", size=9)
        else:
            original = row_cells[0].text.strip()
            set_cell_text(row_cells[0], f"{original}  |  Cumple: X", size=8)

    # Autoevaluacion.
    t = doc.tables[10]
    scores = ["3", "8", "6", "3"]
    for idx, score in enumerate(scores, start=1):
        set_cell_text(t.cell(idx, 3), score, bold=True, size=9)
    set_cell_text(t.cell(5, 3), "20", bold=True, size=9)

    doc.save(str(OUT_DOCX))
    patch_textboxes(OUT_DOCX)
    finalize_textboxes_with_word(OUT_DOCX)


def _textbox_paragraph(text):
    w_ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    p = etree.Element(f"{{{w_ns}}}p")
    p_pr = etree.SubElement(p, f"{{{w_ns}}}pPr")
    spacing = etree.SubElement(p_pr, f"{{{w_ns}}}spacing")
    spacing.set(f"{{{w_ns}}}after", "80")
    r = etree.SubElement(p, f"{{{w_ns}}}r")
    r_pr = etree.SubElement(r, f"{{{w_ns}}}rPr")
    sz = etree.SubElement(r_pr, f"{{{w_ns}}}sz")
    sz.set(f"{{{w_ns}}}val", "18")
    t = etree.SubElement(r, f"{{{w_ns}}}t")
    t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    t.text = text
    return p


def patch_textboxes(docx_path):
    problem_lines = [
        "InnovVentas presenta baja interaccion en su sitio web y una alta tasa de abandono de carrito.",
        "La causa principal es la falta de respuestas rapidas sobre especificaciones, pagos, stock, envios y soporte.",
        "Esta friccion reduce la conversion, afecta la retencion y genera perdidas de ventas en el canal digital.",
    ]
    solution_lines = [
        "Se propone implementar un chatbot 24/7 integrado al e-commerce para responder FAQ, asistir la compra y brindar soporte basico.",
        "La solucion usara Microsoft Copilot Studio, Azure AI Bot Service/Web Chat, Direct Line y Application Insights.",
        "Como evidencias se incluyen respuestas guia, cronograma, recursos, flujo conversacional, arquitectura Azure, pasos de implementacion, metricas y evaluacion.",
    ]
    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    tmp = docx_path.with_suffix(".tmp.docx")
    with zipfile.ZipFile(docx_path, "r") as zin, zipfile.ZipFile(tmp, "w", zipfile.ZIP_DEFLATED) as zout:
        for item in zin.infolist():
            data = zin.read(item.filename)
            if item.filename == "word/document.xml":
                root = etree.fromstring(data)
                boxes = root.xpath(".//w:txbxContent", namespaces=ns)
                empty = [box for box in boxes if not "".join(box.xpath(".//w:t/text()", namespaces=ns)).strip()]
                for idx, box in enumerate(empty[:4]):
                    for child in list(box):
                        box.remove(child)
                    lines = problem_lines if idx < 2 else solution_lines
                    for line in lines:
                        box.append(_textbox_paragraph(line))
                data = etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone=True)
            zout.writestr(item, data)
    tmp.replace(docx_path)


def finalize_textboxes_with_word(docx_path):
    problem = (
        "InnovVentas presenta baja interacción en su sitio web y una alta tasa de abandono de carrito. "
        "La causa principal es la falta de respuestas rápidas sobre especificaciones, pagos, stock, envíos y soporte. "
        "Esta fricción reduce la conversión, afecta la retención y genera pérdidas de ventas en el canal digital."
    )
    solution = (
        "Se propone implementar un chatbot 24/7 integrado al e-commerce para responder FAQ, asistir la compra y brindar soporte básico. "
        "La solución usará Microsoft Copilot Studio, Azure AI Bot Service/Web Chat, Direct Line y Application Insights. "
        "Como evidencias se incluyen respuestas guía, cronograma, recursos, flujo conversacional, arquitectura Azure, pasos de implementación, métricas y evaluación."
    )
    ps = f"""
$docx = '{str(docx_path).replace("'", "''")}'
$word = New-Object -ComObject Word.Application
$word.Visible = $false
try {{
  $doc = $word.Documents.Open($docx)
  $texts = @(
    @'
{problem}
'@,
    @'
{solution}
'@
  )
  for ($idx=0; $idx -lt 2; $idx++) {{
    $s = $doc.Shapes.Item($idx + 2)
    $s.Height = 150
    $s.Width = 452
    $s.TextFrame.MarginLeft = 8
    $s.TextFrame.MarginRight = 8
    $s.TextFrame.MarginTop = 8
    $s.TextFrame.MarginBottom = 8
    $s.TextFrame.TextRange.Text = $texts[$idx]
    $s.TextFrame.TextRange.Font.Name = 'Arial'
    $s.TextFrame.TextRange.Font.Size = 9
    $s.TextFrame.TextRange.Font.Color = 0
  }}
  $doc.Save()
  $doc.Close($false)
}} finally {{
  $word.Quit()
}}
"""
    try:
        subprocess.run(
            ["powershell", "-NoProfile", "-ExecutionPolicy", "Bypass", "-Command", ps],
            check=True,
            capture_output=True,
            text=True,
            timeout=120,
        )
    except Exception as exc:
        print(f"WARNING: no se pudieron finalizar las cajas con Word COM: {exc}")


def write_markdown():
    OUT_MD.parent.mkdir(exist_ok=True)
    lines = [
        "# Propuesta de solucion: Chatbot para InnovVentas",
        "",
        "## Plan de trabajo",
        "",
        "1. Analizar el caso practico y detectar problemas del flujo e-commerce.",
        "2. Levantar preguntas frecuentes e intenciones principales.",
        "3. Disenar el mapa conversacional del chatbot.",
        "4. Seleccionar plataformas Azure y justificar la arquitectura.",
        "5. Definir implementacion, seguridad, metricas y mejora continua.",
        "6. Preparar el entregable final en formato Word.",
        "",
        "## Respuestas a preguntas guia",
        "",
    ]
    for question, answer in GUIDE_QA:
        lines.extend([f"### {question}", "", answer, ""])
    lines.extend(["## Propuesta tecnica", ""])
    for title, paragraphs in PROPOSAL_SECTIONS:
        lines.extend([f"### {title}", ""])
        for paragraph in paragraphs:
            lines.extend([paragraph, ""])
    lines.extend(
        [
            "## Mapa conversacional",
            "",
            "```mermaid",
            "flowchart TD",
            "  A[Cliente entra al sitio] --> B[Chatbot saluda y ofrece opciones]",
            "  B --> C{Intencion detectada}",
            "  C --> D[FAQ: stock, pagos, envios, garantia]",
            "  C --> E[Asistente de compra y recomendacion]",
            "  C --> F[Soporte tecnico basico]",
            "  D --> G[Respuesta y enlace util]",
            "  E --> G",
            "  F --> H{Se resuelve?}",
            "  H -->|Si| G",
            "  H -->|No| I[Escalar a asesor humano]",
            "  G --> J[CSAT y registro de metricas]",
            "  I --> J",
            "```",
            "",
            "## Arquitectura propuesta",
            "",
            "```mermaid",
            "flowchart LR",
            "  U[Usuario web] --> W[Sitio e-commerce + Web Chat]",
            "  W --> D[Direct Line / Canal web]",
            "  D --> B[Copilot Studio o Azure AI Bot Service]",
            "  B --> K[Base de conocimiento FAQ]",
            "  B --> P[API catalogo y stock]",
            "  B --> O[API pedidos y pagos]",
            "  B --> S[Soporte humano]",
            "  B --> M[Application Insights / Azure Monitor]",
            "```",
            "",
            "## Fuentes",
            "",
        ]
    )
    for name, url in AZURE_SOURCES:
        lines.append(f"- [{name}]({url})")
    OUT_MD.write_text("\n".join(lines) + "\n", encoding="utf-8")


def main():
    make_diagrams()
    fill_template()
    write_markdown()
    print(f"DOCX: {OUT_DOCX}")
    print(f"Markdown: {OUT_MD}")


if __name__ == "__main__":
    main()
